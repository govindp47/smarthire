"""
Text extraction service: Extract text from PDF and DOCX files.
"""
import io
import logging
from typing import Optional

from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Extract text from resume files."""
    
    @staticmethod
    async def extract_text_from_bytes(file_content: bytes, file_type: str) -> str:
        """
        Extract text from file content (bytes).
        Works with files from any storage (local, S3, etc.)
        
        Args:
            file_content: File content as bytes
            file_type: File extension (pdf or docx)
            
        Returns:
            Extracted text as string
            
        Raises:
            Exception: If extraction fails
        """
        try:
            if file_type.lower() == "pdf":
                return await TextExtractionService._extract_from_pdf_bytes(file_content)
            elif file_type.lower() in ["docx", "doc"]:
                return await TextExtractionService._extract_from_docx_bytes(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Failed to extract text from {file_type} file: {str(e)}")
            raise
    
    @staticmethod
    async def _extract_from_pdf_bytes(file_content: bytes) -> str:
        """
        Extract text from PDF bytes.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            # Create a BytesIO object from bytes
            pdf_file = io.BytesIO(file_content)
            
            # Read PDF
            reader = PdfReader(pdf_file)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n".join(text_parts)
            
            # Clean up excessive whitespace
            full_text = " ".join(full_text.split())
            
            logger.info(f"Extracted {len(full_text)} characters from PDF ({len(reader.pages)} pages)")
            return full_text
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    async def _extract_from_docx_bytes(file_content: bytes) -> str:
        """
        Extract text from DOCX bytes.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            # Create a BytesIO object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Read DOCX
            doc = docx.Document(docx_file)
            text_parts = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            
            # Clean up excessive whitespace
            full_text = " ".join(full_text.split())
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> list[str]:
        """
        Split text into chunks for processing.
        
        Args:
            text: Full text to chunk
            chunk_size: Target size of each chunk (in characters)
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence end
                sentence_end = text.rfind(". ", start, end)
                if sentence_end != -1 and sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    space = text.rfind(" ", start, end)
                    if space != -1 and space > start + chunk_size // 2:
                        end = space
            
            chunks.append(text[start:end].strip())
            start = end - overlap
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks


# Global instance
text_extractor = TextExtractionService()